from flask import Flask, render_template, request
from werkzeug.utils import secure_filename
import os
from docx import Document
from docx.shared import Cm
import math

UPLOAD_FOLDER = 'static/uploads/'
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = set(['png', 'jpg', 'jpeg'])

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/')
def home():
    return render_template('./index.html')

@app.route('/generate-document', methods=["POST"])
def generate_document():
    # print(request.form)
    form_data = request.form
    number_of_sections = form_data.get('section_count')
    document = Document()
    for i in range(1, int(number_of_sections)+1):
      name_of_defect = form_data.get(f"input_{i}")
      type_of_defect = form_data.get(f"select_{i}")
      desc_of_defect = form_data.get(f"textarea_{i}")
      image_of_defect = request.files[f"image_{i}[]"]
      filename = secure_filename(image_of_defect.filename)
      # image_of_defect.save(os.path.join(app.config['UPLOAD_FOLDER'], filename))

      # Set font
      style = document.styles['Normal']
      style.font.name = 'Calibri'

      # Add paragraphs
      p = document.add_paragraph(f"{i}. Location | Group")
      p.bold = True
      document.add_paragraph(name_of_defect)
      p = document.add_paragraph("Defect Images")
      p.bold = True
      files = request.files.getlist(f"image_{i}[]")
      rows = int(math.ceil(len(files)/2))
      table = document.add_table(rows=rows, cols=2)
      is_odd_images = False
      if rows != len(files) // 2:
         is_odd_images = True
      image_count = 1
      for file in files:
          if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filename_extension = filename.split(".")[1]
            file.save(os.path.join(app.config['UPLOAD_FOLDER'], f"image_{i}_number_{image_count}.png"))
            image_count += 1

      for row in range(rows):
        if row + 1 == rows and is_odd_images:
          left_cell = table.rows[row].cells[0]
          paragraph = left_cell.paragraphs[0]
          left_run = paragraph.add_run()
          left_run.add_picture(f"{UPLOAD_FOLDER}image_{i}_number_{(row+1)*2-1}.png", Cm(7.43))
        else:
          left_cell = table.rows[row].cells[0]
          right_cell = table.rows[row].cells[1]

          paragraph = left_cell.paragraphs[0]
          left_run = paragraph.add_run()
          left_run.add_picture(f"{UPLOAD_FOLDER}image_{i}_number_{(row+1)*2-1}.png", Cm(7.43))
          paragraph = right_cell.paragraphs[0]
          right_run = paragraph.add_run()
          right_run.add_picture(f"{UPLOAD_FOLDER}image_{i}_number_{(row+1)*2-1}.png", Cm(7.43))

      p = document.add_paragraph("Defect Type")
      p.bold = True
      document.add_paragraph(type_of_defect)

      p = document.add_paragraph("Defect Description")
      p.bold = True
      document.add_paragraph(desc_of_defect)
      document.add_page_break()

    document.save('demo.docx')

    return render_template('./report-generated.html')


  # document = Document()

  # # Set font
  # style = document.styles['Normal']
  # style.font.name = 'Calibri'

  # # Add header
  # header_section = document.sections[0]
  # header = header_section.header
  # header_text = header.paragraphs[0]
  # header_text.text = "Header of document"


  # # Add title
  # document.add_heading('This is the Document Title', 0)

  # # Add paragraphs
  # p = document.add_paragraph('A plain paragraph having some ')
  # p.add_run('bold').bold = True
  # p.add_run(' and some ')
  # p.add_run('italic.').italic = True

  # document.add_heading('Heading, level 1', level=1)

  # document.add_paragraph('Intense quote', style='Intense Quote')

  # document.add_paragraph('first item in unordered list', style='List Bullet')

  # document.add_paragraph('first item in ordered list', style='List Number')

  # # Add image
  # document.add_picture('pyword.png', width=Cm(5))

  # records = (
  #     (3, '1', 'Subscribe'),
  #     (7, '12', 'To Python 360'),
  #     (4, '123', 'If you like this video!')
  # )

  # table = document.add_table(rows=1, cols=3)

  # hdr_cells = table.rows[0].cells
  # hdr_cells[0].text = 'Qty'
  # hdr_cells[1].text = 'Id'
  # hdr_cells[2].text = 'Desc'

  # for qty, idx, desc in records:
  #     row_cells = table.add_row().cells
  #     row_cells[0].text = str(qty)
  #     row_cells[1].text = idx
  #     row_cells[2].text = desc

  # document.add_page_break()

  # document.save('demo.docx')

  # # from docx2pdf import convert
  # # docx_file = 'demo.docx'
  # # pdf_file = 'demo.pdf'

  # # convert(docx_file, pdf_file)

if __name__ == '__main__':
    app.run(debug=True)